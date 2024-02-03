import docx
from pathlib import Path
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

my_doc = docx.Document()
my_doc.add_paragraph('This is the first paragraph of a MS Word file. I am very proud of completing this course.')
my_doc.add_paragraph('This is the second paragraph of a MS Word file. I am improving a lot in Python.')

# what if we wanted to add a text in arabic from right to left
paragraph = my_doc.add_paragraph('هذه هي الفقرة الثالثة.')
paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  # this will print the text from right to left

# what if you want the text to go after each other
third_paragraph = my_doc.add_paragraph('This is the third paragraph.')
# add_run() makes the text go after each other
third_paragraph.add_run(' This is a section at the end of the third paragraph')

# add header
my_doc.add_heading('This is level 1 heading', 0)
my_doc.add_heading('This is level 2 heading', 1)
my_doc.add_heading('This is level 3 heading', 2)
my_doc.add_heading('This is level 4 heading', 3)
my_doc.add_heading('This is level 5 heading', 4)
my_doc.add_heading('This is level 6 heading', 5)

# now let's add some style
print(my_doc.paragraphs[0].style)  # checks for the style of the text
print(my_doc.paragraphs[5].style)  # this will print the .style or .text
print(my_doc.paragraphs[4].style)

# now how would you change the style of the text
my_doc.paragraphs[0].style = my_doc.styles['Heading 1']
my_doc.paragraphs[3].style = my_doc.styles['Heading 4']

my_doc.paragraphs[3].style.delete()  # to delete a style

my_doc.add_paragraph('This is a styled paragraph.', 'Title')

# how can we add image to the Word file
my_doc.add_picture(str(Path.home() / Path("OneDrive", "Escritorio", "test.jpg")), width=docx.shared.Inches(5), height=docx.shared.Inches(7))


my_doc.save(Path.home() / Path("OneDrive", "Escritorio", "write.docx"))
